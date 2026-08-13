#!/usr/bin/env python3
"""Generate a realistic Red Herring Prospectus DOCX with embedded PII for testing."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_para(doc, text, bold=False, italic=False, font_size=11, alignment=None):
    """Add a paragraph with optional formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    if alignment:
        para.alignment = alignment
    return para


def add_mixed_para(doc, parts):
    """Add a paragraph with mixed formatting. parts = list of (text, bold, italic)."""
    para = doc.add_paragraph()
    for text, bold, italic in parts:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return para


def create_prospectus():
    doc = Document()

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    add_para(doc, "", font_size=6)
    add_para(doc, "RED HERRING PROSPECTUS", bold=True, font_size=20,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Dated: August 10, 2025", font_size=12,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "(Please read Section 32 of the Companies Act, 2013)",
             italic=True, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "This Red Herring Prospectus will be updated upon filing with the RoC.",
             italic=True, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "", font_size=12)
    add_para(doc, "NEXUS FINTECH SOLUTIONS LIMITED", bold=True, font_size=16,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "(Originally incorporated as 'Nexus Digital Payments Private Limited' on "
             "March 14, 2011 under the Companies Act, 1956. Converted to a Public Limited "
             "Company on June 22, 2023. Certificate of Incorporation issued by the Registrar "
             "of Companies, Mumbai, Maharashtra.)", font_size=9,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "", font_size=6)
    add_para(doc,
             "Registered Office: 4th Floor, Zenith Tower, Plot No. 42, Bandra Kurla Complex, "
             "Bandra East, Mumbai – 400051, Maharashtra, India",
             font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
             "Corporate Office: Unit 1201-1205, Prestige Meridian, MG Road, "
             "Bengaluru – 560001, Karnataka, India",
             font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
             "Tel: +91 22 6789 4321 | Email: ipo@nexusfintech.co.in | "
             "Website: www.nexusfintech.co.in",
             font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
             "Contact Person: Ms. Priya Sharma, Company Secretary and Compliance Officer | "
             "Email: priya.sharma@nexusfintech.co.in | Phone: +91 9876543210",
             font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "", font_size=6)
    add_para(doc, "INITIAL PUBLIC OFFERING OF 2,85,00,000 EQUITY SHARES OF FACE VALUE ₹10 EACH",
             bold=True, font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
             "Price Band: ₹540 – ₹570 per Equity Share (including a premium of ₹530 – ₹560 per share)",
             font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    add_heading_styled(doc, "TABLE OF CONTENTS", level=1)
    toc_items = [
        "1. Definitions and Abbreviations",
        "2. Risk Factors",
        "3. Introduction – Summary of Business",
        "4. Board of Directors and Key Managerial Personnel",
        "5. Promoters and Promoter Group",
        "6. Capital Structure",
        "7. Objects of the Offer",
        "8. Financial Statements",
        "9. Legal and Other Information",
        "10. Statutory and Other Information",
        "11. Offer Information",
    ]
    for item in toc_items:
        add_para(doc, item, font_size=11)

    doc.add_page_break()

    # =========================================================================
    # SECTION 1: DEFINITIONS
    # =========================================================================
    add_heading_styled(doc, "1. DEFINITIONS AND ABBREVIATIONS", level=1)
    add_para(doc,
             "Unless the context otherwise requires, the following terms have the meanings "
             "assigned below in this Red Herring Prospectus:")
    add_para(doc, '"Company" or "Issuer" refers to Nexus Fintech Solutions Limited.')
    add_para(doc, '"BRLMs" refers to the Book Running Lead Managers to the Offer.')
    add_para(doc, '"DRHP" refers to the Draft Red Herring Prospectus dated May 15, 2025.')
    add_para(doc, '"Offer" means the initial public offering of 2,85,00,000 Equity Shares.')
    add_para(doc, '"Registrar" refers to KFin Technologies Limited, the Registrar to the Offer.')
    add_para(doc, '"SEBI" refers to the Securities and Exchange Board of India.')

    doc.add_page_break()

    # =========================================================================
    # SECTION 2: RISK FACTORS
    # =========================================================================
    add_heading_styled(doc, "2. RISK FACTORS", level=1)
    add_para(doc,
             "An investment in equity shares involves a high degree of risk. Prospective "
             "investors should carefully consider all information in this Red Herring Prospectus, "
             "including the risks and uncertainties described below, before making an investment "
             "decision. The following risk factors are material to the Company's business:")

    add_heading_styled(doc, "Internal Risk Factors", level=2)
    add_para(doc,
             "1. Our business is dependent on the continued services of our Promoter and "
             "Managing Director, Mr. Rajesh Kumar Agarwal. The loss of Mr. Agarwal's services "
             "could materially and adversely affect our operations. Mr. Agarwal can be reached "
             "at rajesh.agarwal@nexusfintech.co.in or +91 9823456789 for investor queries "
             "related to the promoter group.")
    add_para(doc,
             "2. Our Chief Technology Officer, Ms. Ananya Iyer, oversees all platform security. "
             "A key-person insurance policy has been taken for Ms. Iyer. Her employee records "
             "indicate Date of Birth: 15/08/1987.")
    add_para(doc,
             "3. Our CFO, Mr. Vikram Singh Rathore, manages all financial operations. "
             "Mr. Rathore joined the company on January 5, 2018. His contact details are: "
             "vikram.rathore@nexusfintech.co.in, Mobile: +91-9834567890.")
    add_para(doc,
             "4. We process an average of ₹2,450 crore in digital transactions per quarter "
             "through our payment gateway. Any disruption to our technology platform, including "
             "from our primary data center at IP 172.16.254.1, could result in significant "
             "revenue loss.")
    add_para(doc,
             "5. The Company's UPI transaction processing system handled 45,00,00,000 "
             "transactions in FY2024-25, with a total GMV of ₹18,750 crore.")

    add_heading_styled(doc, "External Risk Factors", level=2)
    add_para(doc,
             "6. Changes in the regulatory framework by the Reserve Bank of India (RBI) or "
             "SEBI may adversely affect our operations.")
    add_para(doc,
             "7. The fintech industry in India is highly competitive. Competitors include "
             "PhonePe Private Limited, Razorpay Software Private Limited, and BharatPe. "
             "Our market share was approximately 4.2% as of March 31, 2025.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 3: INTRODUCTION
    # =========================================================================
    add_heading_styled(doc, "3. INTRODUCTION – SUMMARY OF BUSINESS", level=1)
    add_para(doc,
             "Nexus Fintech Solutions Limited is a technology-driven financial services "
             "company headquartered in Mumbai, India. The Company provides digital payment "
             "solutions, lending technology, and wealth management platforms to over 1.2 crore "
             "registered users and 85,000 merchant partners across India.")
    add_para(doc,
             "The Company was originally incorporated as Nexus Digital Payments Private Limited "
             "on March 14, 2011 by Mr. Rajesh Kumar Agarwal and Mrs. Sunita Devi Agarwal at "
             "the Registrar of Companies, Mumbai. The Company's CIN is U72200MH2011PLC123456.")

    add_heading_styled(doc, "Corporate Information", level=2)

    # Corporate info table
    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    data = [
        ("CIN", "U72200MH2011PLC123456"),
        ("Registered Office",
         "4th Floor, Zenith Tower, Plot No. 42, Bandra Kurla Complex, "
         "Bandra East, Mumbai – 400051, Maharashtra"),
        ("Corporate Office",
         "Unit 1201-1205, Prestige Meridian, MG Road, "
         "Bengaluru – 560001, Karnataka"),
        ("Telephone", "+91 22 6789 4321"),
        ("Email", "ipo@nexusfintech.co.in"),
        ("Website", "www.nexusfintech.co.in"),
        ("Registrar of Companies", "RoC, Mumbai, Maharashtra"),
        ("Company Secretary", "Ms. Priya Sharma"),
    ]
    for i, (key, val) in enumerate(data):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

    doc.add_page_break()

    # =========================================================================
    # SECTION 4: BOARD OF DIRECTORS AND KMP
    # =========================================================================
    add_heading_styled(doc, "4. BOARD OF DIRECTORS AND KEY MANAGERIAL PERSONNEL", level=1)
    add_para(doc,
             "The following table sets out details of the Directors and Key Managerial "
             "Personnel of the Company as of the date of this Red Herring Prospectus:")

    # Directors table
    add_heading_styled(doc, "4.1 Board of Directors", level=2)

    directors = [
        {
            "name": "Mr. Rajesh Kumar Agarwal",
            "designation": "Chairman & Managing Director",
            "din": "DIN: 00123456",
            "dob": "12/04/1975",
            "address": "B-1402, Oberoi Splendor, Jogeshwari East, Mumbai – 400060, Maharashtra",
            "email": "rajesh.agarwal@nexusfintech.co.in",
            "phone": "+91 9823456789",
            "pan": "ABCPA1234F",
        },
        {
            "name": "Mrs. Sunita Devi Agarwal",
            "designation": "Non-Executive Director",
            "din": "DIN: 00234567",
            "dob": "28/11/1978",
            "address": "B-1402, Oberoi Splendor, Jogeshwari East, Mumbai – 400060, Maharashtra",
            "email": "sunita.agarwal@gmail.com",
            "phone": "+91 9867543210",
            "pan": "BCDPA2345G",
        },
        {
            "name": "Mr. Arjun Mehta",
            "designation": "Independent Director",
            "din": "DIN: 00345678",
            "dob": "05/06/1968",
            "address": "Flat 302, Hiranandani Gardens, Powai, Mumbai – 400076, Maharashtra",
            "email": "arjun.mehta@outlook.com",
            "phone": "+91 9712345678",
            "pan": "CDEPM3456H",
        },
        {
            "name": "Dr. Lakshmi Venkataraman",
            "designation": "Independent Director",
            "din": "DIN: 00456789",
            "dob": "19/02/1972",
            "address": "No. 15, 3rd Cross, Indiranagar, Bengaluru – 560038, Karnataka",
            "email": "lakshmi.venkat@yahoo.co.in",
            "phone": "+91 9845678901",
            "pan": "DEFPV4567J",
        },
        {
            "name": "Mr. Sanjay Prakash Joshi",
            "designation": "Nominee Director (Sequoia Capital)",
            "din": "DIN: 00567890",
            "dob": "30/09/1980",
            "address": "Tower B, Apt 2201, DLF Magnolias, Sector 42, Gurugram – 122002, Haryana",
            "email": "sanjay.joshi@sequoiacap.com",
            "phone": "+91 9988776655",
            "pan": "EFGPJ5678K",
        },
    ]

    table_dir = doc.add_table(rows=len(directors) + 1, cols=6)
    table_dir.style = "Table Grid"
    headers = ["Name", "Designation", "DIN", "Date of Birth", "Address", "Email / Phone"]
    for j, h in enumerate(headers):
        cell = table_dir.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for i, d in enumerate(directors):
        row = table_dir.rows[i + 1]
        row.cells[0].text = d["name"]
        row.cells[1].text = d["designation"]
        row.cells[2].text = d["din"]
        row.cells[3].text = d["dob"]
        row.cells[4].text = d["address"]
        row.cells[5].text = f"{d['email']}\n{d['phone']}"

    add_para(doc, "")

    # =========================================================================
    # KMP Section
    # =========================================================================
    add_heading_styled(doc, "4.2 Key Managerial Personnel", level=2)

    kmp_list = [
        {
            "name": "Ms. Ananya Iyer",
            "designation": "Chief Technology Officer",
            "dob": "15/08/1987",
            "email": "ananya.iyer@nexusfintech.co.in",
            "phone": "+91 9876012345",
            "address": "204, Prestige Lakeside Habitat, Whitefield, Bengaluru – 560066, Karnataka",
        },
        {
            "name": "Mr. Vikram Singh Rathore",
            "designation": "Chief Financial Officer",
            "dob": "22/03/1982",
            "email": "vikram.rathore@nexusfintech.co.in",
            "phone": "+91-9834567890",
            "address": "A-504, Lodha Park, Worli, Mumbai – 400018, Maharashtra",
        },
        {
            "name": "Ms. Priya Sharma",
            "designation": "Company Secretary & Compliance Officer",
            "dob": "07/12/1990",
            "email": "priya.sharma@nexusfintech.co.in",
            "phone": "+91 9876543210",
            "address": "Flat 601, Raheja Atlantis, Parel, Mumbai – 400012, Maharashtra",
        },
        {
            "name": "Mr. Deepak Nair",
            "designation": "Chief Business Officer",
            "dob": "18/01/1985",
            "email": "deepak.nair@nexusfintech.co.in",
            "phone": "+91 8801234567",
            "address": "House No. 47, Sector 15, Chandigarh – 160015, Punjab",
        },
    ]

    for kmp in kmp_list:
        add_para(doc, f"{kmp['name']} – {kmp['designation']}", bold=True, font_size=11)
        add_para(doc,
                 f"Date of Birth: {kmp['dob']}\n"
                 f"Residential Address: {kmp['address']}\n"
                 f"Email: {kmp['email']}\n"
                 f"Contact Number: {kmp['phone']}")

    doc.add_page_break()

    # =========================================================================
    # SECTION 5: PROMOTERS
    # =========================================================================
    add_heading_styled(doc, "5. PROMOTERS AND PROMOTER GROUP", level=1)

    add_heading_styled(doc, "5.1 Details of the Promoters", level=2)

    add_para(doc, "Promoter 1: Mr. Rajesh Kumar Agarwal", bold=True)
    add_para(doc,
             "Mr. Rajesh Kumar Agarwal, aged 50 years, is the founder and Managing Director "
             "of Nexus Fintech Solutions Limited. He holds a B.Tech from IIT Bombay (1997) and "
             "an MBA from IIM Ahmedabad (1999).")
    add_para(doc,
             "Permanent Account Number (PAN): ABCPA1234F\n"
             "Aadhaar Number: 2345 6789 0123\n"
             "Passport Number: J8765432\n"
             "Voter ID: MH/01/042/123456\n"
             "Date of Birth: 12/04/1975\n"
             "Father's Name: Late Mr. Shyam Sunder Agarwal\n"
             "Mother's Name: Mrs. Kamla Devi Agarwal\n"
             "Spouse: Mrs. Sunita Devi Agarwal\n"
             "Residential Address: B-1402, Oberoi Splendor, Jogeshwari East, Mumbai – 400060, Maharashtra\n"
             "Email: rajesh.agarwal@nexusfintech.co.in\n"
             "Phone: +91 9823456789\n"
             "SSN (for US tax reporting purposes): 312-45-6789")

    add_para(doc, "")

    add_para(doc, "Promoter 2: Mrs. Sunita Devi Agarwal", bold=True)
    add_para(doc,
             "Mrs. Sunita Devi Agarwal, aged 47 years, is a Non-Executive Director. She holds "
             "a Bachelor's degree in Commerce from Mumbai University (2000).")
    add_para(doc,
             "Permanent Account Number (PAN): BCDPA2345G\n"
             "Aadhaar Number: 3456 7890 1234\n"
             "Date of Birth: 28/11/1978\n"
             "Father's Name: Mr. Ramesh Chandra Gupta\n"
             "Mother's Name: Mrs. Savitri Devi Gupta\n"
             "Residential Address: B-1402, Oberoi Splendor, Jogeshwari East, Mumbai – 400060, Maharashtra\n"
             "Email: sunita.agarwal@gmail.com\n"
             "Phone: +91 9867543210")

    add_heading_styled(doc, "5.2 Promoter Group Entities", level=2)
    add_para(doc,
             "The following entities form part of the Promoter Group:\n\n"
             "1. Agarwal Family Trust (established on April 10, 2015)\n"
             "   Registered Address: B-1402, Oberoi Splendor, Jogeshwari East, Mumbai – 400060\n\n"
             "2. Nexus Ventures Capital LLP\n"
             "   Registered Address: 301, Marathon Futurex, NM Joshi Marg, Lower Parel, Mumbai – 400013\n"
             "   Email: admin@nexusventures.in\n"
             "   Contact: +91 22 4567 8901\n\n"
             "3. RKA Holdings Private Limited\n"
             "   CIN: U65100MH2015PTC267890\n"
             "   Registered Address: 509, Mittal Tower, Nariman Point, Mumbai – 400021\n"
             "   Email: info@rkaholdings.com\n"
             "   Phone: +91 22 2345 6789")

    doc.add_page_break()

    # =========================================================================
    # SECTION 6: CAPITAL STRUCTURE
    # =========================================================================
    add_heading_styled(doc, "6. CAPITAL STRUCTURE", level=1)
    add_para(doc,
             "The share capital of the Company as of the date of this Red Herring Prospectus "
             "is set out below:")

    cap_table = doc.add_table(rows=5, cols=3)
    cap_table.style = "Table Grid"
    cap_data = [
        ("Particulars", "No. of Shares", "Amount (₹)"),
        ("Authorized Share Capital", "15,00,00,000", "150,00,00,000"),
        ("Issued, Subscribed & Paid-up (Pre-Offer)", "10,00,00,000", "100,00,00,000"),
        ("Present Offer", "2,85,00,000", "28,50,00,000"),
        ("Issued, Subscribed & Paid-up (Post-Offer)", "12,85,00,000", "128,50,00,000"),
    ]
    for i, row_data in enumerate(cap_data):
        for j, val in enumerate(row_data):
            cap_table.rows[i].cells[j].text = val
            if i == 0:
                for run in cap_table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True

    add_para(doc, "")
    add_para(doc, "Shareholding Pattern (Pre-Offer):", bold=True)

    sh_table = doc.add_table(rows=6, cols=4)
    sh_table.style = "Table Grid"
    sh_data = [
        ("Shareholder", "No. of Shares", "% Holding", "Category"),
        ("Mr. Rajesh Kumar Agarwal", "4,50,00,000", "45.00%", "Promoter"),
        ("Mrs. Sunita Devi Agarwal", "1,50,00,000", "15.00%", "Promoter"),
        ("Agarwal Family Trust", "1,00,00,000", "10.00%", "Promoter Group"),
        ("Sequoia Capital India LLP", "2,00,00,000", "20.00%", "Investor"),
        ("ESOP Trust", "1,00,00,000", "10.00%", "Employee Trust"),
    ]
    for i, row_data in enumerate(sh_data):
        for j, val in enumerate(row_data):
            sh_table.rows[i].cells[j].text = val

    doc.add_page_break()

    # =========================================================================
    # SECTION 7: OBJECTS OF THE OFFER
    # =========================================================================
    add_heading_styled(doc, "7. OBJECTS OF THE OFFER", level=1)
    add_para(doc,
             "The Objects of the Offer are:\n\n"
             "1. Funding the expansion of our technology platform and data center infrastructure "
             "(estimated cost: ₹450 crore). Our current primary data center is hosted at IP "
             "address 10.0.128.5 with a disaster recovery site at 192.168.100.25.\n\n"
             "2. Repayment/prepayment of certain borrowings (₹350 crore)\n\n"
             "3. Funding inorganic growth through acquisitions (₹250 crore)\n\n"
             "4. General corporate purposes (₹175.50 crore)\n\n"
             "The main objects clause and objects incidental or ancillary to the main objects of "
             "the Memorandum of Association of the Company enable the Company to undertake its "
             "existing activities and the activities for which funds are being raised through "
             "this Offer.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 8: FINANCIAL STATEMENTS (Summary)
    # =========================================================================
    add_heading_styled(doc, "8. FINANCIAL STATEMENTS (SUMMARY)", level=1)
    add_para(doc,
             "The following table presents a summary of restated consolidated financial "
             "information for Nexus Fintech Solutions Limited for the periods indicated. "
             "The complete audited financial statements are available in Annexure A.")

    fin_table = doc.add_table(rows=8, cols=4)
    fin_table.style = "Table Grid"
    fin_data = [
        ("Particulars (₹ in Crore)", "FY 2024-25", "FY 2023-24", "FY 2022-23"),
        ("Revenue from Operations", "2,845.67", "1,987.23", "1,245.89"),
        ("Other Income", "45.32", "32.18", "21.45"),
        ("Total Expenses", "2,198.45", "1,623.78", "1,089.34"),
        ("Profit Before Tax", "692.54", "395.63", "178.00"),
        ("Tax Expense", "174.23", "99.56", "44.78"),
        ("Profit After Tax", "518.31", "296.07", "133.22"),
        ("Earnings Per Share (₹)", "51.83", "29.61", "13.32"),
    ]
    for i, row_data in enumerate(fin_data):
        for j, val in enumerate(row_data):
            fin_table.rows[i].cells[j].text = val
            if i == 0:
                for run in fin_table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True

    add_para(doc, "")
    add_para(doc, "Statutory Auditors:", bold=True)
    add_para(doc,
             "M/s. Deloitte Haskins & Sells LLP, Chartered Accountants\n"
             "Firm Registration Number: 117366W/W-100018\n"
             "Address: One International Center, Tower 3, 27th Floor, Senapati Bapat Marg, "
             "Elphinstone Road, Mumbai – 400013\n"
             "Partner: Mr. Amit Prakash Kothari (Membership No. 987654)\n"
             "Email: amit.kothari@deloitte.com\n"
             "Phone: +91 22 6185 4000")

    doc.add_page_break()

    # =========================================================================
    # SECTION 9: LEGAL AND OTHER INFORMATION
    # =========================================================================
    add_heading_styled(doc, "9. LEGAL AND OTHER INFORMATION", level=1)

    add_heading_styled(doc, "9.1 Material Litigation", level=2)
    add_para(doc,
             "As of the date of this Red Herring Prospectus, the following material litigation "
             "is pending against the Company:")
    add_para(doc,
             "1. Case No. NCLT/MUM/2024/00789 – Filed by M/s. TechServe Solutions Pvt. Ltd. "
             "claiming ₹12.5 crore for breach of a technology licensing agreement dated "
             "September 15, 2022. The Company has filed its defence and the matter is listed "
             "for hearing on November 20, 2025.")
    add_para(doc,
             "2. Consumer complaint No. CC/2024/3456 filed by Mr. Rohan Dey "
             "(rohan.dey@gmail.com, Contact: +91 9765432109) alleging unauthorized deduction "
             "of ₹45,000 from his account linked to credit card number 4532-0123-4567-8901. "
             "The Company has responded that the transaction was authorized via OTP "
             "verification from IP address 203.0.113.42.")

    add_heading_styled(doc, "9.2 Government Approvals", level=2)
    add_para(doc,
             "The Company has received all necessary approvals from:\n"
             "- Reserve Bank of India (RBI) – Payment Aggregator License dated February 28, 2023\n"
             "- SEBI – Registration as Investment Adviser (INA000012345) dated July 10, 2022\n"
             "- Ministry of Corporate Affairs – CIN: U72200MH2011PLC123456")

    doc.add_page_break()

    # =========================================================================
    # SECTION 10: STATUTORY AND OTHER INFORMATION
    # =========================================================================
    add_heading_styled(doc, "10. STATUTORY AND OTHER INFORMATION", level=1)

    add_heading_styled(doc, "10.1 Book Running Lead Managers", level=2)
    add_para(doc,
             "1. Kotak Mahindra Capital Company Limited\n"
             "   Address: 27BKC, C-27, G Block, Bandra Kurla Complex, Bandra East, "
             "Mumbai – 400051\n"
             "   Contact Person: Mr. Rashi Patil\n"
             "   Email: rashi.patil@kotak.com\n"
             "   Phone: +91 22 4336 0000\n"
             "   SEBI Registration: INM000008704")
    add_para(doc, "")
    add_para(doc,
             "2. ICICI Securities Limited\n"
             "   Address: ICICI Centre, H.T. Parekh Marg, Churchgate, Mumbai – 400020\n"
             "   Contact Person: Ms. Neha Kapoor\n"
             "   Email: neha.kapoor@icicisecurities.com\n"
             "   Phone: +91 22 2288 2460\n"
             "   SEBI Registration: INM000011179")

    add_heading_styled(doc, "10.2 Registrar to the Offer", level=2)
    add_para(doc,
             "KFin Technologies Limited\n"
             "Selenium, Tower B, Plot No. 31-32, Financial District, Nanakramguda, "
             "Hyderabad – 500032, Telangana\n"
             "Contact Person: Mr. Suresh Babu Ponnada\n"
             "Email: suresh.ponnada@kfintech.com\n"
             "Phone: +91 40 6716 2222\n"
             "Website: www.kfintech.com\n"
             "SEBI Registration: INR000000221")

    add_heading_styled(doc, "10.3 Bankers to the Offer", level=2)

    add_para(doc, "Escrow Collection Bank:", bold=True)
    add_para(doc,
             "HDFC Bank Limited\n"
             "Branch: Fort Branch, Mumbai\n"
             "Account Name: Nexus Fintech IPO – Escrow Account\n"
             "Account Number: 50200087654321\n"
             "IFSC Code: HDFC0000060\n"
             "Contact: Mr. Amit Jain, Email: amit.jain@hdfcbank.com, Phone: +91 9920123456")

    add_para(doc, "Refund Bank:", bold=True)
    add_para(doc,
             "ICICI Bank Limited\n"
             "Branch: BKC Branch, Mumbai\n"
             "Account Name: Nexus Fintech IPO – Refund Account\n"
             "Account Number: 123409876543\n"
             "IFSC Code: ICIC0000004\n"
             "Contact: Ms. Kavita Reddy, Email: kavita.reddy@icicibank.com, Phone: +91 9876012345")

    doc.add_page_break()

    # =========================================================================
    # SECTION 10.4: SYNDICATE MEMBERS AND ADDITIONAL CONTACTS
    # =========================================================================
    add_heading_styled(doc, "10.4 Experts and Other Intermediaries", level=2)

    add_para(doc, "Legal Advisor to the Company:", bold=True)
    add_para(doc,
             "Cyril Amarchand Mangaldas\n"
             "Peninsula Chambers, Peninsula Corporate Park, Lower Parel, Mumbai – 400013\n"
             "Partner: Ms. Divya Srinivasan\n"
             "Email: divya.srinivasan@cyrilshroff.com\n"
             "Phone: +91 22 2496 4455")

    add_para(doc, "Legal Advisor to the BRLMs:", bold=True)
    add_para(doc,
             "AZB & Partners\n"
             "AZB House, Peninsula Corporate Park, Lower Parel, Mumbai – 400013\n"
             "Partner: Mr. Karan Bhasin\n"
             "Email: karan.bhasin@azbpartners.com\n"
             "Phone: +91 22 6639 6880")

    add_para(doc, "Intellectual Property Advisor:", bold=True)
    add_para(doc,
             "Remfry & Sagar\n"
             "Millennium Plaza, Sector 27, Gurugram – 122009, Haryana\n"
             "Contact: Dr. Pooja Mathur\n"
             "Email: pooja.mathur@remfry.com\n"
             "Phone: +91 124 280 6100")

    doc.add_page_break()

    # =========================================================================
    # SECTION 11: OFFER INFORMATION
    # =========================================================================
    add_heading_styled(doc, "11. OFFER INFORMATION", level=1)

    add_heading_styled(doc, "11.1 Offer Details", level=2)

    offer_table = doc.add_table(rows=8, cols=2)
    offer_table.style = "Table Grid"
    offer_data = [
        ("Offer Opens", "September 15, 2025"),
        ("Offer Closes", "September 18, 2025"),
        ("Price Band", "₹540 – ₹570 per Share"),
        ("Bid Lot", "25 Equity Shares"),
        ("Minimum Bid Amount", "₹14,250 (at upper price band)"),
        ("Face Value", "₹10 per Equity Share"),
        ("Listing", "BSE Limited and National Stock Exchange of India Limited"),
        ("Offer Size", "₹1,624.50 crore (at upper price band)"),
    ]
    for i, (key, val) in enumerate(offer_data):
        offer_table.rows[i].cells[0].text = key
        offer_table.rows[i].cells[1].text = val

    add_para(doc, "")

    add_heading_styled(doc, "11.2 Allocation Details", level=2)
    add_para(doc,
             "The Offer is being made in terms of Rule 19(2)(b) of the Securities Contracts "
             "(Regulation) Rules, 1957.\n\n"
             "Category-wise allocation:\n"
             "• Qualified Institutional Buyers (QIBs): Not more than 50%\n"
             "• Non-Institutional Investors: Not less than 15%\n"
             "• Retail Individual Investors: Not less than 35%")

    add_heading_styled(doc, "11.3 Grading", level=2)
    add_para(doc,
             "This Offer has been graded by CRISIL Limited as 'CRISIL IPO Grade 4/5' "
             "indicating 'Above Average Fundamentals'. The grading report dated July 25, 2025 "
             "is available on the website of CRISIL at www.crisil.com.")

    doc.add_page_break()

    # =========================================================================
    # ADDITIONAL SECTION: EMPLOYEE DETAILS (for more PII)
    # =========================================================================
    add_heading_styled(doc, "ANNEXURE B: KEY EMPLOYEE DETAILS (CONFIDENTIAL)", level=1)
    add_para(doc,
             "The following information is provided for regulatory compliance and is "
             "confidential in nature.", italic=True)

    employees = [
        {
            "name": "Rashi Patil",
            "designation": "Vice President – Investment Banking",
            "email": "rashi.patil@gmail.com",
            "phone": "+91 9876543210",
            "dob": "March 22, 1992",
            "address": "Flat 12B, Sea View Apartments, Carter Road, Bandra West, Mumbai – 400050",
            "ssn": "456-78-9012",
            "cc": "4532 8901 2345 6789",
        },
        {
            "name": "Rohan Dey",
            "designation": "Senior Manager – Technology",
            "email": "rohan.dey@gmail.com",
            "phone": "+91 9765432109",
            "dob": "July 14, 1988",
            "address": "House 23, Salt Lake City, Sector V, Kolkata – 700091, West Bengal",
            "ssn": "567-89-0123",
            "cc": "5425 6789 0123 4567",
        },
        {
            "name": "Meera Krishnamurthy",
            "designation": "Head of Compliance",
            "email": "meera.krishnamurthy@nexusfintech.co.in",
            "phone": "+91 8765432109",
            "dob": "November 3, 1985",
            "address": "No. 8, 2nd Main Road, Koramangala 4th Block, Bengaluru – 560034, Karnataka",
            "ssn": "678-90-1234",
            "cc": "3782 822463 10005",
        },
        {
            "name": "Aditya Prakash Verma",
            "designation": "Lead Data Scientist",
            "email": "aditya.verma@nexusfintech.co.in",
            "phone": "91-7890123456",
            "dob": "February 9, 1991",
            "address": "C-302, Jaypee Greens, Sector 128, Noida – 201304, Uttar Pradesh",
            "ssn": "789-01-2345",
            "cc": "6011 1234 5678 9012",
        },
    ]

    emp_table = doc.add_table(rows=len(employees) + 1, cols=7)
    emp_table.style = "Table Grid"
    emp_headers = ["Name", "Designation", "DOB", "Email", "Phone", "SSN (US Tax)", "Card on File"]
    for j, h in enumerate(emp_headers):
        emp_table.rows[0].cells[j].text = h
        for run in emp_table.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)

    for i, emp in enumerate(employees):
        row = emp_table.rows[i + 1]
        row.cells[0].text = emp["name"]
        row.cells[1].text = emp["designation"]
        row.cells[2].text = emp["dob"]
        row.cells[3].text = emp["email"]
        row.cells[4].text = emp["phone"]
        row.cells[5].text = emp["ssn"]
        row.cells[6].text = emp["cc"]

    add_para(doc, "")

    # More inline PII references
    add_heading_styled(doc, "Travel and Expense Records", level=2)
    add_para(doc,
             "Recent business travel expenses submitted by key personnel:\n\n"
             "1. Mr. Rajesh Kumar Agarwal – Flight booking to New York (September 5-12, 2025) "
             "charged to corporate credit card ending 8901. Full card number on file: "
             "4111 1111 1111 1111. Booking confirmation sent to rajesh.agarwal@nexusfintech.co.in.\n\n"
             "2. Ms. Ananya Iyer – Conference registration for AWS re:Invent 2025. "
             "Payment made via card number 5500-0000-0000-0004. Receipt sent to "
             "ananya.iyer@nexusfintech.co.in from IP 198.51.100.73.\n\n"
             "3. Mr. Deepak Nair – Hotel booking in Singapore (October 1-5, 2025). "
             "Charged to personal card 3400 000000 00009. Passport Number: K1234567. "
             "Confirmation email forwarded from deepak.nair@nexusfintech.co.in.")

    doc.add_page_break()

    # =========================================================================
    # DECLARATION PAGE
    # =========================================================================
    add_heading_styled(doc, "DECLARATION", level=1)
    add_para(doc,
             "We, the undersigned Directors of Nexus Fintech Solutions Limited, hereby declare "
             "that all statements made in this Red Herring Prospectus are true and correct to "
             "the best of our knowledge and belief, and that there are no other facts, the "
             "omission of which would make any statement in this Red Herring Prospectus "
             "misleading or untrue.")

    add_para(doc, "")
    add_para(doc, "Signed by all Directors:", bold=True)
    add_para(doc, "")

    signatories = [
        ("Mr. Rajesh Kumar Agarwal", "Chairman & Managing Director"),
        ("Mrs. Sunita Devi Agarwal", "Non-Executive Director"),
        ("Mr. Arjun Mehta", "Independent Director"),
        ("Dr. Lakshmi Venkataraman", "Independent Director"),
        ("Mr. Sanjay Prakash Joshi", "Nominee Director"),
    ]

    for name, title in signatories:
        add_para(doc, "")
        add_para(doc, "________________________")
        add_para(doc, f"{name}", bold=True)
        add_para(doc, f"{title}")

    add_para(doc, "")
    add_para(doc, f"Date: August 10, 2025")
    add_para(doc, f"Place: Mumbai")

    add_para(doc, "")
    add_para(doc,
             "For and on behalf of the Board of Directors of Nexus Fintech Solutions Limited\n"
             "Registered Office: 4th Floor, Zenith Tower, Plot No. 42, Bandra Kurla Complex, "
             "Bandra East, Mumbai – 400051, Maharashtra, India\n"
             "CIN: U72200MH2011PLC123456\n"
             "Email: ipo@nexusfintech.co.in\n"
             "Website: www.nexusfintech.co.in",
             font_size=9)

    # Save
    output_path = "Red_Herring_Prospectus.docx"
    doc.save(output_path)
    print(f"✅ Red Herring Prospectus generated: {output_path}")
    print(f"   Pages: ~15-18 (estimated)")
    print(f"   PII embedded: names, emails, phones, addresses, SSNs, credit cards, DOBs, IPs, companies")


if __name__ == "__main__":
    create_prospectus()
