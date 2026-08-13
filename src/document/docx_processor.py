import logging
from dataclasses import dataclass
from typing import Tuple, List, Optional
import docx
from docx.text.paragraph import Paragraph

from src.core.entity import PIIEntity

logger = logging.getLogger(__name__)

@dataclass
class TextSegment:
    """Represents a segment of text from a DOCX file and its origin."""
    text: str
    start_offset: int
    end_offset: int
    source_type: str  # 'paragraph', 'table', 'header', 'footer'
    source_index: int
    paragraph_index: Optional[int] = None
    cell_coords: Optional[Tuple[int, int]] = None


class DocxProcessor:
    """Handles reading, extracting text, and applying redactions to DOCX files."""
    
    def __init__(self):
        pass

    def extract_text(self, filepath: str) -> str:
        """Extracts all text from a DOCX file."""
        text, _ = self.extract_text_with_positions(filepath)
        return text

    def extract_text_with_positions(self, filepath: str) -> Tuple[str, List[TextSegment]]:
        """
        Extracts text from a DOCX file and tracks the origin of each segment.
        Returns the full text and a list of TextSegment objects mapping offsets.
        """
        try:
            document = docx.Document(filepath)
        except Exception as e:
            logger.error(f"Failed to open DOCX file {filepath}: {e}")
            raise

        segments = []
        full_text = ""
        current_offset = 0

        # Extract from Paragraphs
        for i, para in enumerate(document.paragraphs):
            text = para.text
            if text:
                segments.append(TextSegment(
                    text=text,
                    start_offset=current_offset,
                    end_offset=current_offset + len(text),
                    source_type='paragraph',
                    source_index=i,
                    paragraph_index=i
                ))
                full_text += text + "\n"
                current_offset += len(text) + 1
            else:
                full_text += "\n"
                current_offset += 1

        # Extract from Tables
        for i, table in enumerate(document.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        text = para.text
                        if text:
                            segments.append(TextSegment(
                                text=text,
                                start_offset=current_offset,
                                end_offset=current_offset + len(text),
                                source_type='table',
                                source_index=i,
                                paragraph_index=p_idx,
                                cell_coords=(r_idx, c_idx)
                            ))
                            full_text += text + "\n"
                            current_offset += len(text) + 1
                        else:
                            full_text += "\n"
                            current_offset += 1

        # Extract from Sections (Headers and Footers)
        for i, section in enumerate(document.sections):
            for p_idx, para in enumerate(section.header.paragraphs):
                text = para.text
                if text:
                    segments.append(TextSegment(
                        text=text,
                        start_offset=current_offset,
                        end_offset=current_offset + len(text),
                        source_type='header',
                        source_index=i,
                        paragraph_index=p_idx
                    ))
                    full_text += text + "\n"
                    current_offset += len(text) + 1
                else:
                    full_text += "\n"
                    current_offset += 1

            for p_idx, para in enumerate(section.footer.paragraphs):
                text = para.text
                if text:
                    segments.append(TextSegment(
                        text=text,
                        start_offset=current_offset,
                        end_offset=current_offset + len(text),
                        source_type='footer',
                        source_index=i,
                        paragraph_index=p_idx
                    ))
                    full_text += text + "\n"
                    current_offset += len(text) + 1
                else:
                    full_text += "\n"
                    current_offset += 1

        return full_text, segments

    def apply_redactions(self, input_path: str, output_path: str, entities: List[PIIEntity]) -> None:
        """
        Applies redactions to a DOCX file in-place, preserving formatting.
        Saves the redacted version to output_path.
        """
        if not entities:
            logger.info("No entities to redact. Saving original document as-is.")
            document = docx.Document(input_path)
            document.save(output_path)
            return

        document = docx.Document(input_path)
        _, segments = self.extract_text_with_positions(input_path)
        
        # Group entities by the text segment they belong to
        segment_entities = {id(seg): [] for seg in segments}
        for entity in entities:
            for seg in segments:
                # Check if entity falls within this segment
                if seg.start_offset <= entity.start < seg.end_offset:
                    rel_start = entity.start - seg.start_offset
                    rel_end = entity.end - seg.start_offset
                    segment_entities[id(seg)].append((entity, rel_start, rel_end))
                    break
                    
        for seg in segments:
            entities_in_seg = segment_entities[id(seg)]
            if not entities_in_seg:
                continue
                
            para = None
            if seg.source_type == 'paragraph':
                para = document.paragraphs[seg.source_index]
            elif seg.source_type == 'table':
                para = document.tables[seg.source_index].rows[seg.cell_coords[0]].cells[seg.cell_coords[1]].paragraphs[seg.paragraph_index]
            elif seg.source_type == 'header':
                para = document.sections[seg.source_index].header.paragraphs[seg.paragraph_index]
            elif seg.source_type == 'footer':
                para = document.sections[seg.source_index].footer.paragraphs[seg.paragraph_index]
                
            if para is not None:
                self._replace_in_paragraph(para, entities_in_seg)

        document.save(output_path)
        logger.info(f"Saved redacted document to {output_path}")
        
    def _replace_in_paragraph(self, paragraph: Paragraph, entities_in_para: List[Tuple[PIIEntity, int, int]]) -> None:
        """
        Replaces text within a paragraph's runs while preserving formatting.
        Handles cases where entities span across multiple runs.
        """
        if not paragraph.runs:
            return

        # Sort in reverse order to process from end of paragraph to start,
        # preventing offset shifts for earlier entities.
        entities_in_para.sort(key=lambda x: x[1], reverse=True)
        
        # Extract working copies of run texts to manipulate them before assigning back
        run_texts = [run.text for run in paragraph.runs]
        
        # Build a mapping of character index (relative to full para string) to (run_index, char_index_in_run)
        char_to_run = []
        for run_idx, run_text in enumerate(run_texts):
            for char_idx in range(len(run_text)):
                char_to_run.append((run_idx, char_idx))
                
        for entity, rel_start, rel_end in entities_in_para:
            if rel_start >= len(char_to_run) or rel_end > len(char_to_run) or rel_start >= rel_end:
                logger.warning(f"Invalid bounds for entity '{entity.text}' in paragraph. Max len: {len(char_to_run)}")
                continue
                
            start_run_idx, start_char_idx = char_to_run[rel_start]
            end_run_idx, end_char_idx = char_to_run[rel_end - 1]
            
            replacement = entity.replacement or "[REDACTED]"
            
            if start_run_idx == end_run_idx:
                # Entity is completely within a single run
                text = run_texts[start_run_idx]
                run_texts[start_run_idx] = text[:start_char_idx] + replacement + text[end_char_idx + 1:]
            else:
                # Entity spans multiple runs
                start_text = run_texts[start_run_idx]
                run_texts[start_run_idx] = start_text[:start_char_idx] + replacement
                
                # Clear intermediate runs
                for i in range(start_run_idx + 1, end_run_idx):
                    run_texts[i] = ""
                    
                end_text = run_texts[end_run_idx]
                run_texts[end_run_idx] = end_text[end_char_idx + 1:]
                
        # Assign the modified texts back to the original runs, preserving formatting
        for run_idx, run in enumerate(paragraph.runs):
            run.text = run_texts[run_idx]
